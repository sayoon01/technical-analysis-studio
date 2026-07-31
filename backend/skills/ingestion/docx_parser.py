"""DOCX ingestion (paragraphs as text blocks)."""

from __future__ import annotations

from pathlib import Path

from backend.skills.ingestion.pdf_parser import RawPage, RawTextBlock
from backend.skills.ingestion.text_parser import TextDocument, _split_by_size


def parse_docx(path: str | Path) -> TextDocument:
    try:
        from docx import Document
    except ImportError as e:
        raise RuntimeError(
            "python-docx is required for DOCX ingestion. pip install python-docx"
        ) from e

    doc = Document(str(path))
    paras = [p.text.strip() for p in doc.paragraphs if p.text.strip()]
    # tables as line-joined cells
    for table in doc.tables:
        for row in table.rows:
            cells = [c.text.strip() for c in row.cells if c.text.strip()]
            if cells:
                paras.append(" | ".join(cells))

    full = "\n\n".join(paras)
    chunks = _split_by_size(full, 3000)
    pages: list[RawPage] = []
    for i, chunk in enumerate(chunks, start=1):
        block = RawTextBlock(
            text=chunk,
            bbox=(0.0, 0.0, 612.0, 792.0),
            block_type="TEXT",
            confidence=1.0,
        )
        pages.append(
            RawPage(
                page_number=i,
                width=612.0,
                height=792.0,
                text_layer_available=True,
                full_text=chunk,
                blocks=[block],
            )
        )
    return TextDocument(pages=pages)
