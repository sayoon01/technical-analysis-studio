"""DOCX exporter via python-docx."""

from __future__ import annotations

import re
from pathlib import Path

from docx import Document
from docx.shared import Inches, Pt


def export_docx(
    path: Path,
    markdown: str,
    *,
    title: str,
    image_root: Path | None = None,
) -> Path:
    path.parent.mkdir(parents=True, exist_ok=True)
    doc = Document()
    style = doc.styles["Normal"]
    style.font.name = "NanumGothic"
    style.font.size = Pt(11)

    doc.add_heading(title, level=0)

    for block in markdown.split("\n"):
        line = block.rstrip()
        if not line:
            continue
        if line.startswith("# "):
            # already have title
            continue
        if line.startswith("## "):
            doc.add_heading(line[3:].strip(), level=1)
            continue
        if line.startswith("### "):
            doc.add_heading(line[4:].strip(), level=2)
            continue
        if line.startswith("#### "):
            doc.add_heading(line[5:].strip(), level=3)
            continue
        img = re.match(r"!\[([^\]]*)\]\(([^)]+)\)", line)
        if img and image_root is not None:
            rel = img.group(2)
            img_path = (image_root / Path(rel).name) if "visuals/" in rel else Path(rel)
            if not img_path.is_file():
                img_path = image_root / Path(rel).name
            if img_path.is_file():
                doc.add_picture(str(img_path), width=Inches(5.5))
                cap = img.group(1).strip()
                if cap:
                    doc.add_paragraph(cap)
                continue
        # strip md emphasis lightly
        text = re.sub(r"[*`_]", "", line)
        if text.startswith("|") and text.endswith("|"):
            doc.add_paragraph(text)
        else:
            doc.add_paragraph(text)

    doc.save(str(path))
    return path
